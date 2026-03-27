"""
厦门大学学位论文格式自动化工具
基于《厦门大学管理学院MBA学位论文开题报告格式模板》202511版
"""

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import re
import copy


# ============================================================
# 厦门大学论文格式规范常量
# ============================================================

PAGE_MARGIN_TOP = Cm(2.54)      # 1 inch，官方规范
PAGE_MARGIN_BOTTOM = Cm(2.54)   # 1 inch，官方规范
PAGE_MARGIN_LEFT = Cm(2.8)
PAGE_MARGIN_RIGHT = Cm(2.8)
PAGE_HEADER_DISTANCE = Cm(1.5)
PAGE_FOOTER_DISTANCE = Cm(1.75)

FONT_CHINESE_BODY = '宋体'
FONT_CHINESE_HEADING = '黑体'
FONT_CHINESE_COVER = '楷体'
FONT_ENGLISH = 'Times New Roman'

SIZE_NORMAL = Pt(12)       # 小四
SIZE_H1 = Pt(15)           # 小三
SIZE_H2 = Pt(14)           # 四号
SIZE_H3 = Pt(12)           # 小四
SIZE_H4 = Pt(12)           # 小四
SIZE_HEADER = Pt(10.5)
SIZE_FOOTER = Pt(10.5)
SIZE_CAPTION = Pt(10.5)

FIRST_LINE_INDENT = Pt(24)  # 2个12pt字符 = 首行缩进2字


# 一级标题识别模式（第X章）
HEADING1_PATTERNS = [
    r'^第\s*[一二三四五六七八九十百\d]+\s*章',
    r'^Chapter\s+\d+',
    r'^CHAPTER\s+\d+',
]

# 二级标题识别模式（第X节）
HEADING2_PATTERNS = [
    r'^第\s*[一二三四五六七八九十百\d]+\s*节',
    r'^Section\s+\d+',
]

# 三级标题识别模式
HEADING3_PATTERNS = [
    r'^[一二三四五六七八九十]+[、．.\s]',
    r'^\d+\.\d+\s+\S',
]

# 四级标题识别模式
HEADING4_PATTERNS = [
    r'^\([一二三四五六七八九十\d]+\)',
    r'^\d+\.\d+\.\d+\s+\S',
]

# 特殊段落识别
ABSTRACT_ZH_PATTERNS = [r'^摘\s*要$', r'^摘　要$']
ABSTRACT_EN_PATTERNS = [r'^Abstract$', r'^ABSTRACT$']
KEYWORDS_PATTERNS = [r'^关键词[：:]', r'^Keywords[：:]']
REFERENCE_PATTERNS = [r'^参考文献[：:]?\s*$', r'^References[：:]?\s*$']
ACKNOWLEDGMENT_PATTERNS = [r'^致\s*谢$', r'^Acknowledgments?$']

# 特殊节标题（需居中）：摘要、目录、参考文献、致谢等
SECTION_TITLE_PATTERNS = (
    ABSTRACT_ZH_PATTERNS + ABSTRACT_EN_PATTERNS +
    REFERENCE_PATTERNS + ACKNOWLEDGMENT_PATTERNS +
    [r'^目\s*录$', r'^目　录$', r'^Contents?$', r'^TABLE OF CONTENTS$']
)


def _set_style_fonts(style_elem, ascii_font=None, east_asia_font=None, size_pt=None,
                     bold=None, italic=None):
    """修改样式定义中的字体属性"""
    rPr = style_elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style_elem.append(rPr)

    if ascii_font or east_asia_font:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        if ascii_font:
            rFonts.set(qn('w:ascii'), ascii_font)
            rFonts.set(qn('w:hAnsi'), ascii_font)
        if east_asia_font:
            rFonts.set(qn('w:eastAsia'), east_asia_font)
            rFonts.set(qn('w:cs'), east_asia_font)

    if size_pt is not None:
        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz')
            rPr.append(sz)
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        half_pts = str(int(size_pt * 2))
        sz.set(qn('w:val'), half_pts)
        szCs.set(qn('w:val'), half_pts)

    if bold is not None:
        b_elem = rPr.find(qn('w:b'))
        if bold:
            if b_elem is None:
                b_elem = OxmlElement('w:b')
                rPr.append(b_elem)
        else:
            if b_elem is not None:
                rPr.remove(b_elem)

    if italic is not None:
        i_elem = rPr.find(qn('w:i'))
        if italic:
            if i_elem is None:
                i_elem = OxmlElement('w:i')
                rPr.append(i_elem)
        else:
            if i_elem is not None:
                rPr.remove(i_elem)


def _set_style_paragraph_format(style_elem, alignment=None, line_spacing=None,
                                  space_before=None, space_after=None,
                                  first_line_indent=None, left_indent=None):
    """修改样式定义中的段落格式属性"""
    pPr = style_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        style_elem.insert(0, pPr)

    if alignment is not None:
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            pPr.append(jc)
        align_map = {
            WD_ALIGN_PARAGRAPH.LEFT: 'left',
            WD_ALIGN_PARAGRAPH.CENTER: 'center',
            WD_ALIGN_PARAGRAPH.RIGHT: 'right',
            WD_ALIGN_PARAGRAPH.JUSTIFY: 'both',
        }
        jc.set(qn('w:val'), align_map.get(alignment, 'left'))

    if line_spacing is not None:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        if isinstance(line_spacing, float) and line_spacing <= 4.0:
            # 多倍行距
            spacing.set(qn('w:line'), str(int(line_spacing * 240)))
            spacing.set(qn('w:lineRule'), 'auto')
        else:
            # 固定值（EMU转twips）
            twips = int(line_spacing / 914400 * 1440)
            spacing.set(qn('w:line'), str(twips))
            spacing.set(qn('w:lineRule'), 'exact')

    if space_before is not None or space_after is not None:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        if space_before is not None:
            twips = int(space_before / 914400 * 1440)
            spacing.set(qn('w:before'), str(twips))
        if space_after is not None:
            twips = int(space_after / 914400 * 1440)
            spacing.set(qn('w:after'), str(twips))

    if first_line_indent is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        if first_line_indent == 0:
            ind.attrib.pop(qn('w:firstLine'), None)
            ind.set(qn('w:firstLine'), '0')
        else:
            twips = int(first_line_indent / 914400 * 1440)
            ind.set(qn('w:firstLine'), str(twips))

    if left_indent is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        twips = int(left_indent / 914400 * 1440)
        ind.set(qn('w:left'), str(twips))


def _set_toc_tab_stop(style_elem, pos_twips):
    """为目录样式设置右对齐点线引导符 tab stop（清除已有 tabs 后重设）"""
    pPr = style_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        # 插在 rPr 之前
        rPr = style_elem.find(qn('w:rPr'))
        if rPr is not None:
            rPr.addprevious(pPr)
        else:
            style_elem.append(pPr)
    # 清除旧 tabs，重新设置
    old_tabs = pPr.find(qn('w:tabs'))
    if old_tabs is not None:
        pPr.remove(old_tabs)
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(pos_twips))
    tab.set(qn('w:leader'), 'dot')
    tabs.append(tab)
    pPr.append(tabs)


def update_styles(doc):
    """更新文档中的样式定义以符合厦大格式"""
    changes = []

    style_configs = [
        # (样式名, ascii字体, 中文字体, 字号pt, 加粗, 对齐, 行距倍数, 段前pt, 段后pt, 首行缩进)
        ('Normal', FONT_ENGLISH, FONT_CHINESE_BODY, 12, False,
         WD_ALIGN_PARAGRAPH.JUSTIFY, 1.5, None, None, FIRST_LINE_INDENT),
        ('Heading 1', FONT_ENGLISH, FONT_CHINESE_HEADING, 15, True,
         WD_ALIGN_PARAGRAPH.CENTER, 1.0, Pt(12), Pt(5), Pt(0)),
        ('Heading 2', FONT_ENGLISH, FONT_CHINESE_HEADING, 14, True,
         WD_ALIGN_PARAGRAPH.CENTER, 1.0, Pt(5), Pt(2.5), Pt(0)),
        ('Heading 3', FONT_ENGLISH, FONT_CHINESE_HEADING, 12, True,
         WD_ALIGN_PARAGRAPH.LEFT, 1.5, Pt(2.5), None, Pt(0)),
        ('Heading 4', FONT_ENGLISH, FONT_CHINESE_HEADING, 12, True,
         WD_ALIGN_PARAGRAPH.LEFT, 1.5, None, None, Pt(0)),
    ]

    for (style_name, ascii_font, zh_font, size_pt, bold,
         align, line_sp, sp_before, sp_after, first_indent) in style_configs:
        try:
            style = doc.styles[style_name]
            elem = style.element
            _set_style_fonts(elem, ascii_font=ascii_font, east_asia_font=zh_font,
                             size_pt=size_pt, bold=bold)
            _set_style_paragraph_format(
                elem,
                alignment=align,
                line_spacing=line_sp,
                space_before=sp_before,
                space_after=sp_after,
                first_line_indent=first_indent,
            )
            changes.append(f'更新样式：{style_name}（{zh_font} {size_pt}pt）')
        except KeyError:
            changes.append(f'未找到样式：{style_name}，已跳过')

    # 同步更新所有通过 outlineLvl 检测到的标题样式（如 标题 1/2/3/4）
    # 确保中文或自定义命名的标题样式获得正确的字体、字号、对齐设置
    _heading_lvl_config = {
        1: (15, True,  WD_ALIGN_PARAGRAPH.CENTER, 1.0, Pt(12),  Pt(5),   Pt(0)),
        2: (14, True,  WD_ALIGN_PARAGRAPH.CENTER, 1.0, Pt(5),   Pt(2.5), Pt(0)),
        3: (12, True,  WD_ALIGN_PARAGRAPH.LEFT,   1.5, Pt(2.5), None,    Pt(0)),
        4: (12, True,  WD_ALIGN_PARAGRAPH.LEFT,   1.5, None,    None,    Pt(0)),
    }
    _standard_heading_names = {f'Heading {n}' for n in range(1, 9)}
    _doc_heading_styles = _get_heading_styles(doc)
    for _sname, _lvl in _doc_heading_styles.items():
        if _sname in _standard_heading_names:
            continue  # 已在上面的循环中处理
        if _lvl not in _heading_lvl_config:
            continue
        _size_pt, _bold, _align, _line_sp, _sp_before, _sp_after, _first_indent = _heading_lvl_config[_lvl]
        try:
            _style = doc.styles[_sname]
            _set_style_fonts(_style.element, ascii_font=FONT_ENGLISH,
                             east_asia_font=FONT_CHINESE_HEADING, size_pt=_size_pt, bold=_bold)
            _set_style_paragraph_format(_style.element,
                                         alignment=_align,
                                         line_spacing=_line_sp,
                                         space_before=_sp_before,
                                         space_after=_sp_after,
                                         first_line_indent=_first_indent)
            changes.append(f'更新自定义标题样式：{_sname}（黑体 {_size_pt}pt 级别{_lvl}）')
        except KeyError:
            pass

    # 摘要/目录标题样式：黑体三号（16pt）加粗居中，1.5倍行距
    # 若样式不存在则从 Normal 派生创建
    for style_name in ('摘要目录标题',):
        try:
            style = doc.styles[style_name]
            created = False
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles['Normal']
            created = True
        _set_style_fonts(style.element, ascii_font=FONT_ENGLISH,
                         east_asia_font=FONT_CHINESE_HEADING, size_pt=16, bold=True)
        _set_style_paragraph_format(style.element,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                     line_spacing=1.5,
                                     first_line_indent=Pt(0))
        verb = '创建' if created else '更新'
        changes.append(f'{verb}样式：摘要目录标题（黑体 16pt 三号 居中）')

    # toc 1/2/3 样式：目录条目，若不存在则创建
    # A4，左右各2.8cm：文本宽度 = 15.4cm = 8731 twips，用于右对齐页码 tab stop
    TOC_TAB_POS = 8731  # twips

    # 目录层级缩进（twips）：toc1=0，toc2=2字符，toc3=4字符（1字符=12pt*20=240twips）
    TOC_LEVEL_INDENT = {1: 0, 2: 480, 3: 960}

    # toc 1: 14pt 加粗；toc 2: 12pt 加粗；toc 3: 12pt 非加粗（参考范文模板）
    toc_configs = [
        ('toc 1', 1, 14, True),
        ('toc 2', 2, 12, True),
        ('toc 3', 3, 12, False),
    ]
    for (tname, level, toc_size, toc_bold) in toc_configs:
        left_twips = TOC_LEVEL_INDENT[level]
        # left_indent 以 EMU 传入：twips → EMU = twips / 1440 * 914400
        left_emu = int(left_twips / 1440 * 914400)
        try:
            tstyle = doc.styles[tname]
        except KeyError:
            tstyle = doc.styles.add_style(tname, WD_STYLE_TYPE.PARAGRAPH)
            tstyle.base_style = doc.styles['Normal']
        _set_style_fonts(tstyle.element, ascii_font=FONT_ENGLISH,
                         east_asia_font=FONT_CHINESE_BODY, size_pt=toc_size, bold=toc_bold)
        _set_style_paragraph_format(tstyle.element,
                                     alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                     line_spacing=1.5,
                                     first_line_indent=Pt(0),
                                     left_indent=left_emu)
        # 清除样式级 numPr（列表编号会引入额外缩进）和悬挂缩进
        pPr = tstyle.element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                ind.attrib.pop(qn('w:hanging'), None)
        _set_toc_tab_stop(tstyle.element, TOC_TAB_POS)

    # 同步更新中文本地化目录样式（目录 1/2/3）
    # toc 1/目录 1: 14pt 加粗；toc 2/目录 2: 12pt 加粗；toc 3/目录 3: 12pt 非加粗
    ZH_TOC_NAMES = {
        '目录 1': 1, '目录 2': 2, '目录 3': 3,
        '目录1':  1, '目录2':  2, '目录3':  3,
    }
    _ZH_TOC_SIZE = {1: 14, 2: 12, 3: 12}
    _ZH_TOC_BOLD = {1: True, 2: True, 3: False}
    for zname, level in ZH_TOC_NAMES.items():
        try:
            zstyle = doc.styles[zname]
        except KeyError:
            continue
        left_twips = TOC_LEVEL_INDENT[level]
        left_emu = int(left_twips / 1440 * 914400)
        _set_style_fonts(zstyle.element, ascii_font=FONT_ENGLISH,
                         east_asia_font=FONT_CHINESE_BODY,
                         size_pt=_ZH_TOC_SIZE[level], bold=_ZH_TOC_BOLD[level])
        _set_style_paragraph_format(zstyle.element,
                                     alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                     line_spacing=1.5,
                                     first_line_indent=Pt(0),
                                     left_indent=left_emu)
        pPr = zstyle.element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                ind.attrib.pop(qn('w:hanging'), None)
        _set_toc_tab_stop(zstyle.element, TOC_TAB_POS)

    # ── 修正所有 toc 相关样式定义（不依赖精确名称，扫描所有段落样式）──
    # 解决 doc.styles['toc 1'] 在中文 Word 中可能失败（样式名为"目录 1"/"TOC 1"）的问题
    _toc_level_re = re.compile(r'(?:toc|目录)\s*([123])', re.IGNORECASE)
    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        m = _toc_level_re.search(style.name) or \
            _toc_level_re.search(style.element.get(qn('w:styleId'), ''))
        if not m:
            continue
        level = int(m.group(1))
        if level not in (1, 2, 3):
            continue
        # 从样式定义中移除 numPr 和 ind（彻底清除任何继承的缩进）
        s_pPr = style.element.find(qn('w:pPr'))
        if s_pPr is None:
            s_pPr = OxmlElement('w:pPr')
            style.element.insert(0, s_pPr)
        for tag in (qn('w:numPr'),):
            e = s_pPr.find(tag)
            if e is not None:
                s_pPr.remove(e)
        s_ind = s_pPr.find(qn('w:ind'))
        if s_ind is None:
            s_ind = OxmlElement('w:ind')
            s_pPr.append(s_ind)
        for attr in (qn('w:hanging'), qn('w:firstLine')):
            s_ind.attrib.pop(attr, None)
        s_ind.set(qn('w:left'), str(TOC_LEVEL_INDENT.get(level, 0)))
        s_ind.set(qn('w:firstLine'), '0')

    # ── 对中文目录区域内所有段落强制去除左侧缩进 ──
    # 边界："目录"标题段落 → 第一个匹配HEADING1的章节标题段落（不含）
    # 不依赖样式名，覆盖所有本地化变体
    TOC_TITLE_PATS = [r'^目\s*录$', r'^目　录$']
    all_paras = doc.paragraphs
    zh_toc_title_idx = -1
    for i, para in enumerate(all_paras):
        if _matches_any(para.text.strip(), TOC_TITLE_PATS):
            zh_toc_title_idx = i
            break

    if zh_toc_title_idx >= 0:
        # 找中文目录结束位置：遇到一级标题段落或英文目录标题即停止
        _hs = _get_heading_styles(doc)
        zh_toc_end_idx = len(all_paras)
        for i in range(zh_toc_title_idx + 1, len(all_paras)):
            para = all_paras[i]
            if _hs.get(para.style.name) == 1 and para.text.strip():
                zh_toc_end_idx = i
                break
            # 遇到另一个目录标题（Contents）也停止
            if _matches_any(para.text.strip(), [r'^Contents?$', r'^TABLE OF CONTENTS$']):
                zh_toc_end_idx = i
                break

        def _fix_toc_para_indent(para):
            # 从样式名中判断层级，决定缩进量
            sname = para.style.name
            style_id = para.style.element.get(qn('w:styleId'), '')
            toc_m = re.search(r'(?:toc|目录)\s*([123])', sname, re.IGNORECASE) or \
                    re.search(r'(?:toc|目录)\s*([123])', style_id, re.IGNORECASE)
            level = int(toc_m.group(1)) if toc_m else 1
            left_str = str(TOC_LEVEL_INDENT.get(level, 0))

            p_elem = para._p
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p_elem.insert(0, pPr)
            for tag in (qn('w:numPr'), qn('w:tabs')):
                e = pPr.find(tag)
                if e is not None:
                    pPr.remove(e)
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind')
                pPr.append(ind)
            for attr in (qn('w:hanging'), qn('w:firstLine')):
                ind.attrib.pop(attr, None)
            ind.set(qn('w:left'), left_str)
            ind.set(qn('w:firstLine'), '0')
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'left')

        for para in all_paras[zh_toc_title_idx + 1: zh_toc_end_idx]:
            _fix_toc_para_indent(para)

    # 页眉页脚样式：五号宋体（10.5pt）居中，非加粗（规范明确）
    for style_name in ('Header', 'Footer'):
        try:
            style = doc.styles[style_name]
            _set_style_fonts(style.element, ascii_font=FONT_ENGLISH,
                             east_asia_font=FONT_CHINESE_BODY, size_pt=10.5, bold=False)
            _set_style_paragraph_format(style.element,
                                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                         line_spacing=1.0)
        except KeyError:
            pass

    return changes


def update_page_setup(doc):
    """更新所有节的页面设置"""
    changes = []
    for i, section in enumerate(doc.sections):
        section.top_margin = PAGE_MARGIN_TOP
        section.bottom_margin = PAGE_MARGIN_BOTTOM
        section.left_margin = PAGE_MARGIN_LEFT
        section.right_margin = PAGE_MARGIN_RIGHT
        section.header_distance = PAGE_HEADER_DISTANCE
        section.footer_distance = PAGE_FOOTER_DISTANCE
    changes.append('页面设置：A4，上下边距2.54cm，左右边距2.8cm，页眉1.5cm，页脚1.75cm')
    return changes


def _add_bookmark(p_elem, bookmark_id, name):
    """在段落中添加书签（用于 PAGEREF 引用）。
    bookmarkStart 必须在 w:pPr 之后、内容 run 之前，否则 Word 计算页码会出错。
    """
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bookmark_id))
    bm_start.set(qn('w:name'), name)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bookmark_id))
    # 必须放在 w:pPr 之后（OOXML 规范要求）
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is not None:
        pPr.addnext(bm_start)
    elif len(p_elem):
        p_elem[0].addprevious(bm_start)
    else:
        p_elem.append(bm_start)
    p_elem.append(bm_end)


def _get_max_bookmark_id(doc):
    """获取文档中已有书签的最大 ID"""
    max_id = 0
    for bm in doc.element.body.iter(qn('w:bookmarkStart')):
        try:
            bid = int(bm.get(qn('w:id'), 0))
            if bid > max_id:
                max_id = bid
        except (ValueError, TypeError):
            pass
    return max_id


def _make_pageref_runs(bookmark_name):
    """创建 PAGEREF 域 run 列表（显示书签所在页码）"""
    r1 = OxmlElement('w:r')
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    r1.append(fc1)

    r2 = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = f' PAGEREF {bookmark_name} '
    r2.append(instr)

    r3 = OxmlElement('w:r')
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'separate')
    r3.append(fc2)

    r4 = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '?'
    r4.append(t)

    r5 = OxmlElement('w:r')
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'end')
    r5.append(fc3)

    return [r1, r2, r3, r4, r5]


def _make_text_run(text):
    """创建纯文字run"""
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    return r


def _make_page_number_field():
    """创建页码域代码（仅域，不含破折号）"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run1 = OxmlElement('w:r')
    run1.append(copy.deepcopy(fldChar1))
    run2 = OxmlElement('w:r')
    run2.append(instrText)
    run3 = OxmlElement('w:r')
    run3.append(fldChar2)
    return run1, run2, run3


def _set_footer_page_number(footer):
    """
    在页脚中写入居中页码，格式：- X -
    奇数页脚和偶数页脚共用此函数。
    彻底清除页脚 XML 内所有原有内容后重建，避免原文页码残留导致重复。
    """
    try:
        if footer.is_linked_to_previous:
            footer.is_linked_to_previous = False
    except Exception:
        return

    # 获取页脚 XML 根元素
    try:
        ftr_xml = footer._element
    except Exception:
        return
    if ftr_xml is None:
        return

    # 彻底清除所有子元素（段落、文本框、表格、SDT 等），从零重建
    for child in list(ftr_xml):
        ftr_xml.remove(child)

    # 新建段落
    para_elem = OxmlElement('w:p')
    ftr_xml.append(para_elem)

    # 设置段落样式和居中对齐
    pPr = para_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para_elem.insert(0, pPr)
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = OxmlElement('w:pStyle')
        pPr.insert(0, pStyle)
    pStyle.set(qn('w:val'), 'Footer')
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'center')

    # 插入 "- " + PAGE域 + " -"
    def text_run(txt):
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = txt
        r.append(t)
        return r

    para_elem.append(text_run('- '))
    r1, r2, r3 = _make_page_number_field()
    para_elem.append(r1)
    para_elem.append(r2)
    para_elem.append(r3)
    para_elem.append(text_run(' -'))


def update_footer_page_numbers(doc, thesis_title=''):
    """为所有节添加页码，格式：- X -，奇偶页均设置（符合厦大规范）"""
    changes = []
    for section in doc.sections:
        _set_footer_page_number(section.footer)            # 奇数页脚
        _set_footer_page_number(section.even_page_footer)  # 偶数页脚

    changes.append('页脚：已添加自动页码（格式：- X -，奇偶页均设置）')
    return changes


def _matches_any(text, patterns):
    """检查文本是否匹配任意模式"""
    for p in patterns:
        if re.match(p, text.strip()):
            return True
    return False


def _get_heading_styles(doc):
    """
    检测所有通过 w:outlineLvl 定义标题级别的段落样式。
    返回 {style_name: level}，level 为 1-based（1=一级标题，2=二级标题…）。
    与样式名无关，可处理中文名（标题 1）、英文名（Heading 1）及任意自定义名称。
    """
    result = {}
    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        pPr = style.element.find(qn('w:pPr'))
        if pPr is None:
            continue
        ol = pPr.find(qn('w:outlineLvl'))
        if ol is None:
            continue
        val = ol.get(qn('w:val'))
        if val is None:
            continue
        try:
            lvl = int(val)
            if 0 <= lvl <= 8:
                result[style.name] = lvl + 1  # outlineLvl 0 → level 1
        except ValueError:
            pass
    return result


def auto_detect_headings(doc):
    """
    对没有应用标题样式的段落，根据内容模式自动检测并应用标题样式。
    只处理当前样式为 Normal/正文/Body Text 的段落。
    """
    changes = []
    normal_style_names = {'Normal', 'Body Text', 'Body', '正文', '正文-缩进',
                          'Default Paragraph Style', 'No Spacing'}
    heading_applied = 0

    heading_styles = _get_heading_styles(doc)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        current_style = para.style.name
        # 只对非标题样式的段落做检测（使用 outlineLvl 判断）
        if current_style in heading_styles:
            continue

        if current_style not in normal_style_names:
            continue

        if _matches_any(text, HEADING1_PATTERNS):
            para.style = doc.styles['Heading 1']
            heading_applied += 1
        elif _matches_any(text, HEADING2_PATTERNS):
            para.style = doc.styles['Heading 2']
            heading_applied += 1
        elif _matches_any(text, HEADING3_PATTERNS):
            para.style = doc.styles['Heading 3']
            heading_applied += 1
        elif _matches_any(text, HEADING4_PATTERNS):
            para.style = doc.styles['Heading 4']
            heading_applied += 1

    if heading_applied > 0:
        changes.append(f'自动识别标题：共为 {heading_applied} 个段落应用了标题样式')
    return changes


def fix_heading_direct_format(doc):
    """
    清除标题段落上的直接格式覆盖（direct formatting），
    确保标题样式定义生效。
    只处理第一个真实章节标题（第X章）及其之后的段落，
    避免封面/声明页的偶发标题被错误修改。
    保留加粗/斜体等字符级格式不变。
    """
    changes = []
    cleared = 0

    # 使用 outlineLvl 检测所有标题样式（与样式名无关）
    heading_styles = _get_heading_styles(doc)

    # 从目录之后开始处理标题（封面/声明/摘要/目录区域不做任何修改）
    chapter_style = _get_chapter_heading_style(doc)
    post_toc_idx = _find_post_toc_idx(doc)
    # 在目录之后找到第一个章节标题（不依赖 HEADING1_PATTERNS，支持任意命名）
    first_chapter_idx = None
    for i, para in enumerate(doc.paragraphs):
        if i < post_toc_idx:
            continue
        if para.style.name == chapter_style and para.text.strip():
            first_chapter_idx = i
            break

    for i, para in enumerate(doc.paragraphs):
        # 使用 outlineLvl 判断是否为标题（支持 Heading N、标题 N 及任意自定义样式）
        if para.style.name not in heading_styles:
            continue
        # 跳过目录之前的所有段落（封面/声明/摘要/目录）
        if i < post_toc_idx:
            continue
        # 若目录后仍未找到章节，则跳过第一章之前的段落
        if first_chapter_idx is not None and i < first_chapter_idx:
            continue
        # 清除段落级别的直接格式（pPr上的格式覆盖）
        pPr = para._p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            para._p.insert(0, pPr)
        # 清除 spacing（间距）直接覆盖
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            pPr.remove(spacing)
        # 标题首行缩进强制置0：不能仅删除 w:ind，否则会从 Normal 样式继承 firstLine=480
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLine'), '0')
        ind.attrib.pop(qn('w:hanging'), None)

        # 强制对齐：一级/二级标题居中，三级/四级标题左对齐（使用 outlineLvl 级别）
        level = heading_styles.get(para.style.name)
        if level is not None:
            align_val = 'center' if level <= 2 else 'left'
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), align_val)

        # 章节标题（一级标题样式）：强制另起一页
        is_chapter = (para.style.name == chapter_style and bool(para.text.strip()))
        pgBr = pPr.find(qn('w:pageBreakBefore'))
        if is_chapter:
            if pgBr is None:
                pgBr = OxmlElement('w:pageBreakBefore')
                pPr.append(pgBr)
            pgBr.set(qn('w:val'), '1')
        else:
            # 非章节标题：清除可能残留的分页符
            if pgBr is not None:
                pPr.remove(pgBr)

        # 清除run级别的字体/大小覆盖（保留bold/italic）
        for run in para.runs:
            rPr = run._r.find(qn('w:rPr'))
            if rPr is None:
                continue
            for tag in [qn('w:rFonts'), qn('w:sz'), qn('w:szCs'),
                        qn('w:color'), qn('w:highlight')]:
                elem = rPr.find(tag)
                if elem is not None:
                    rPr.remove(elem)
        cleared += 1

    if cleared > 0:
        changes.append(f'清除标题直接格式：共处理 {cleared} 个标题段落')

    # ── 特殊节标题居中（摘要、参考文献、致谢、目录等）──
    # 这些段落通常不使用 Heading 样式，需直接写入居中对齐并清除首行缩进
    section_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if _matches_any(text, SECTION_TITLE_PATTERNS):
            p_elem = para._p
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p_elem.insert(0, pPr)
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'center')
            # 清除首行缩进（节标题不应有首行缩进）
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                ind.attrib.pop(qn('w:firstLine'), None)
                ind.attrib.pop(qn('w:hanging'), None)
            section_count += 1
    if section_count:
        changes.append(f'节标题居中：已居中 {section_count} 个特殊节标题（摘要/参考文献/致谢等）')

    return changes


FIGURE_CAPTION_PATTERNS = [
    r'^图\s*[\d一二三四五六七八九十]',
    r'^Figure\s+\d',
    r'^Fig\.\s*\d',
]
TABLE_CAPTION_PATTERNS = [
    r'^表\s*[\d一二三四五六七八九十]',
    r'^Table\s+\d',
]
ALL_CAPTION_PATTERNS = FIGURE_CAPTION_PATTERNS + TABLE_CAPTION_PATTERNS


def _is_caption_para(para):
    """判断段落是否为题注（按样式名或文字模式）"""
    sname = para.style.name.lower()
    if 'caption' in sname or '题注' in sname:
        return True
    text = para.text.strip()
    return any(re.match(p, text) for p in ALL_CAPTION_PATTERNS)


def _is_figure_caption(para):
    text = para.text.strip()
    return any(re.match(p, text) for p in FIGURE_CAPTION_PATTERNS)


def _is_table_caption(para):
    text = para.text.strip()
    return any(re.match(p, text) for p in TABLE_CAPTION_PATTERNS)


def _has_drawing(para):
    """段落中是否包含嵌入图片"""
    p = para._p
    return (p.find('.//' + qn('w:drawing')) is not None or
            p.find('.//' + qn('w:pict')) is not None)


def _set_pPr_flag(pPr, tag_name):
    """在 pPr 中确保某个布尔标志元素存在（如 w:keepNext / w:keepLines）"""
    if pPr.find(qn(tag_name)) is None:
        elem = OxmlElement(tag_name)
        pPr.append(elem)


def _get_or_create_pPr(para_elem):
    pPr = para_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para_elem.insert(0, pPr)
    return pPr


def fix_captions_and_tables(doc):
    """
    修正题注格式，防止图表跨页：
    1. 更新 Caption/题注 样式定义（宋体 10.5pt 居中 无缩进）
    2. 对所有识别为题注的段落直接应用格式
    3. 图片段落 → keepWithNext（与下方图题注保持同页）
    4. 图题注 → keepLines（不拆行）
    5. 表题注 → keepWithNext（与下方表格保持同页）+ keepLines
    6. 表格行 → cantSplit（禁止单行跨页）
    """
    changes = []

    # ── 1. 更新样式定义 ──────────────────────────────────────────
    for sname in ('Caption', '题注'):
        try:
            style = doc.styles[sname]
            _set_style_fonts(style.element, ascii_font=FONT_ENGLISH,
                             east_asia_font=FONT_CHINESE_BODY, size_pt=10.5, bold=True)
            _set_style_paragraph_format(
                style.element,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                line_spacing=1.5,
                first_line_indent=Pt(0),
            )
            # 清除悬挂缩进
            pPr = style.element.find(qn('w:pPr'))
            if pPr is not None:
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    for attr in (qn('w:hanging'), qn('w:firstLine'), qn('w:left')):
                        ind.attrib.pop(attr, None)
            changes.append(f'更新题注样式：{sname}（宋体 10.5pt 加粗 居中）')
        except KeyError:
            pass

    # ── 2. 直接格式化题注段落 + 图片段落处理 ─────────────────────
    paragraphs = doc.paragraphs
    caption_count = 0
    figure_count = 0

    for i, para in enumerate(paragraphs):
        # 跳过表格内段落，不修改表格单元格中的内容
        if para._p.getparent() is not None and para._p.getparent().tag == qn('w:tc'):
            continue

        is_cap = _is_caption_para(para)
        has_img = _has_drawing(para)

        if not is_cap and not has_img:
            continue

        p_elem = para._p
        pPr = _get_or_create_pPr(p_elem)

        if is_cap:
            caption_count += 1
            # 直接写入字体和对齐（覆盖直接格式）
            # 清除旧 rPr 并设置字体
            for r_elem in p_elem.findall(qn('w:r')):
                rPr = r_elem.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r_elem.insert(0, rPr)
                # 字体
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:ascii'), FONT_ENGLISH)
                rFonts.set(qn('w:hAnsi'), FONT_ENGLISH)
                rFonts.set(qn('w:eastAsia'), FONT_CHINESE_BODY)
                rFonts.set(qn('w:cs'), FONT_CHINESE_BODY)
                # 字号 10.5pt = 21 half-points
                for sz_name in ('w:sz', 'w:szCs'):
                    sz = rPr.find(qn(sz_name))
                    if sz is None:
                        sz = OxmlElement(sz_name)
                        rPr.append(sz)
                    sz.set(qn('w:val'), '21')
                # 题注加粗（参考范文模板）
                for b_name in ('w:b', 'w:bCs'):
                    b_elem = rPr.find(qn(b_name))
                    if b_elem is None:
                        b_elem = OxmlElement(b_name)
                        rPr.append(b_elem)
                    b_elem.set(qn('w:val'), '1')

            # 段落对齐居中
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'center')

            # 清除段落首行缩进
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                for attr in (qn('w:firstLine'), qn('w:hanging'), qn('w:left')):
                    ind.attrib.pop(attr, None)

            # keepLines（题注不拆行）
            _set_pPr_flag(pPr, 'w:keepLines')

            # 表题注加 keepWithNext（下方紧接表格）
            if _is_table_caption(para):
                _set_pPr_flag(pPr, 'w:keepNext')

        if has_img:
            figure_count += 1
            # 图片段落 keepWithNext（下方紧接图题注）
            _set_pPr_flag(pPr, 'w:keepNext')
            _set_pPr_flag(pPr, 'w:keepLines')

    if caption_count:
        changes.append(f'题注格式：已修正 {caption_count} 个题注段落（宋体 10.5pt 加粗 居中）')
    if figure_count:
        changes.append(f'图片：已为 {figure_count} 个图片段落设置 keepWithNext')
    return changes


def fix_body_text_fonts(doc):
    """
    修正正文段落格式：首行缩进2字、1.5倍行距、段前段后0pt、宋体/Times New Roman 12pt、两端对齐。
    仅处理第一章之后的非标题、非目录、非空段落。
    """
    changes = []

    # 使用 outlineLvl 检测所有标题样式（与样式名无关，支持 Heading N/标题 N/自定义名）
    heading_styles = _get_heading_styles(doc)
    _SKIP_PREFIXES = ('toc ', 'TOC ', 'List')
    _SKIP_NAMES = {
        '摘要目录标题', 'Caption', '题注', 'Header', 'Footer',
        'footnote text', 'endnote text', 'footnote reference',
    }

    chapter_style = _get_chapter_heading_style(doc)
    post_toc_idx = _find_post_toc_idx(doc)
    first_chapter_idx = None
    for i, para in enumerate(doc.paragraphs):
        if i < post_toc_idx:
            continue
        if para.style.name == chapter_style and para.text.strip():
            first_chapter_idx = i
            break

    fixed = 0
    in_references = False  # 是否处于参考文献区域
    for i, para in enumerate(doc.paragraphs):
        if i < post_toc_idx:
            continue
        if first_chapter_idx is not None and i < first_chapter_idx:
            continue

        text = para.text.strip()
        sname = para.style.name

        # 检测进入/离开参考文献区域
        if _matches_any(text, REFERENCE_PATTERNS):
            in_references = True
        elif in_references and sname in heading_styles and text:
            # 遇到下一个标题（致谢等）则离开参考文献区域
            in_references = False

        # 跳过标题样式（通过 outlineLvl 检测，与名称无关）
        if sname in heading_styles:
            continue
        # 跳过目录、题注、页眉页脚等特殊样式
        if any(sname.startswith(p) for p in _SKIP_PREFIXES):
            continue
        if sname in _SKIP_NAMES:
            continue
        if 'caption' in sname.lower() or '题注' in sname:
            continue

        # 跳过表格内段落（不修改表格单元格内容）
        if para._p.getparent() is not None and para._p.getparent().tag == qn('w:tc'):
            continue

        # 跳过空段落
        if not text:
            continue

        # 跳过特殊节标题（摘要/参考文献/致谢等），避免覆盖居中对齐和字号
        if _matches_any(text, SECTION_TITLE_PATTERNS):
            continue

        p_elem = para._p
        pPr = p_elem.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_elem.insert(0, pPr)

        # 两端对齐（参考文献条目左对齐顶格）
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            pPr.append(jc)
        jc.set(qn('w:val'), 'left' if in_references else 'both')

        # 行距 1.5 倍，段前/段后 0pt
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:line'), '360')       # 1.5 × 240 = 360
        spacing.set(qn('w:lineRule'), 'auto')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')

        # 缩进：参考文献条目顶格，正文首行缩进 2 字（480 twips）
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        if in_references:
            ind.set(qn('w:firstLine'), '0')
            ind.attrib.pop(qn('w:hanging'), None)
            ind.attrib.pop(qn('w:left'), None)
        else:
            ind.set(qn('w:firstLine'), '480')
            ind.attrib.pop(qn('w:hanging'), None)
            ind.attrib.pop(qn('w:left'), None)

        # 清除段首手动缩进（制表符 / 全角/半角空格），参考文献区域也执行
        first_text_run = None
        for run in para.runs:
            if run.text:
                first_text_run = run
                break
        if first_text_run is not None:
            cleaned = first_text_run.text.lstrip('\t\u3000 ')
            if cleaned != first_text_run.text:
                first_text_run.text = cleaned

        # 修正每个 run 的字体和字号（保留 bold/italic/underline/color 等）
        for run in para.runs:
            if not run.text:
                continue
            r_elem = run._r
            rPr = r_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r_elem.insert(0, rPr)

            # 字体：中文宋体，英文/数字 Times New Roman
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'),     FONT_ENGLISH)
            rFonts.set(qn('w:hAnsi'),     FONT_ENGLISH)
            rFonts.set(qn('w:eastAsia'),  FONT_CHINESE_BODY)
            rFonts.set(qn('w:cs'),        FONT_ENGLISH)

            # 字号 12pt（= 24 halfpoints）
            for sz_name in ('w:sz', 'w:szCs'):
                sz_elem = rPr.find(qn(sz_name))
                if sz_elem is None:
                    sz_elem = OxmlElement(sz_name)
                    rPr.append(sz_elem)
                sz_elem.set(qn('w:val'), '24')

        fixed += 1

    if fixed > 0:
        changes.append(f'修正正文格式：{fixed} 个段落（首行缩进2字、1.5倍行距、宋体/TNR 12pt）')
    return changes


def _make_rPr(ascii_font='Times New Roman', east_asia_font='宋体', size_pt=10.5, bold=False):
    """创建带字体设置的 rPr 元素（用于页眉文字）"""
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), east_asia_font)
    rFonts.set(qn('w:cs'), east_asia_font)  # 复杂文种字体（部分中文字符走此路径）
    rPr.append(rFonts)
    half = str(int(size_pt * 2))
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), half)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), half)
    rPr.append(sz)
    rPr.append(szCs)
    # 明确设置加粗状态（页眉规范为非加粗）
    b = OxmlElement('w:b')
    b.set(qn('w:val'), '1' if bold else '0')
    rPr.append(b)
    bCs = OxmlElement('w:bCs')
    bCs.set(qn('w:val'), '1' if bold else '0')
    rPr.append(bCs)
    return rPr


def _clear_header_content(header):
    """清空页眉中所有段落的内容（保留段落属性 pPr）"""
    for para in header.paragraphs:
        p_elem = para._p
        to_remove = [
            child for child in p_elem
            if (child.tag.split('}')[-1] if '}' in child.tag else child.tag) != 'pPr'
        ]
        for child in to_remove:
            p_elem.remove(child)


def _make_styleref_field(style_name='Heading 1', backward_lookup=True):
    """
    创建带字体的 StyleRef 域，自动显示当前页最近的指定样式文本。
    backward_lookup=True  → 加 \\l 标志（向后查找，适合章节页眉）
    backward_lookup=False → 不加 \\l（向前查找，适合摘要/目录页眉）
    """
    rPr = _make_rPr()

    run_begin = OxmlElement('w:r')
    run_begin.append(copy.deepcopy(rPr))
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    run_begin.append(fc1)

    run_instr = OxmlElement('w:r')
    run_instr.append(copy.deepcopy(rPr))
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    bl_flag = ' \\l' if backward_lookup else ''
    instr.text = f' STYLEREF "{style_name}"{bl_flag} \\* MERGEFORMAT '
    run_instr.append(instr)

    run_sep = OxmlElement('w:r')
    run_sep.append(copy.deepcopy(rPr))
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'separate')
    run_sep.append(fc2)

    run_placeholder = OxmlElement('w:r')
    run_placeholder.append(copy.deepcopy(rPr))
    t = OxmlElement('w:t')
    t.text = '第X章 标题'
    run_placeholder.append(t)

    run_end = OxmlElement('w:r')
    run_end.append(copy.deepcopy(rPr))
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'end')
    run_end.append(fc3)

    return run_begin, run_instr, run_sep, run_placeholder, run_end


def _extract_thesis_title(doc):
    """
    从文档封面自动提取论文标题（中文）。
    按优先级依次尝试：封面标题样式 → Title样式 → 文档前5页中最长的加粗大字段落
    """
    # 优先：封面标题样式
    cover_styles = {'封面标题（中）', '封面标题', 'Title', '标题'}
    for para in doc.paragraphs:
        if para.style.name in cover_styles and para.text.strip():
            return para.text.strip()

    # 次选：文档开头（前30段）中字号最大且加粗的段落
    candidates = []
    for para in doc.paragraphs[:30]:
        text = para.text.strip()
        if not text or len(text) < 4:
            continue
        for run in para.runs:
            size = run.font.size
            bold = run.bold or para.style.font.bold
            if size and bold:
                candidates.append((size, text))
                break

    if candidates:
        candidates.sort(key=lambda x: -x[0])
        return candidates[0][1]

    return ''


def _get_chapter_heading_style(doc):
    """
    检测文档实际使用的章节标题（一级标题）样式名。
    使用 w:outlineLvl 进行检测，与样式名无关（支持 Heading 1、标题 1 等任意名称）。
    优先返回有实际段落使用的最低级别（level=1）样式名。
    """
    heading_styles = _get_heading_styles(doc)

    # 一级标题样式集合（outlineLvl=0 → level=1）
    level1_styles = {name for name, lvl in heading_styles.items() if lvl == 1}

    # 优先：实际在文档中使用的一级标题样式
    for para in doc.paragraphs:
        if para.style.name in level1_styles and para.text.strip():
            return para.style.name

    # 回退：使用任意级别标题样式（取最低级别）
    min_level = None
    min_style = None
    for name, lvl in heading_styles.items():
        for para in doc.paragraphs:
            if para.style.name == name and para.text.strip():
                if min_level is None or lvl < min_level:
                    min_level = lvl
                    min_style = name
                break
    if min_style:
        return min_style

    return 'Heading 1'


def _find_post_toc_idx(doc):
    """
    找到目录结束后第一个正文段落的索引。
    扫描所有段落，找到最后一个目录相关段落（目录标题或 toc 样式段落），
    返回其后一个段落的索引。
    若找不到目录，退而寻找第一个匹配 HEADING1_PATTERNS 的章节标题段落，
    以此作为正文起点（避免封面/声明页的 Heading 样式段落被误处理）。
    若两者均找不到，返回 0。
    """
    TOC_TITLE_PATS = [r'^目\s*录$', r'^目　录$', r'^Contents?$', r'^TABLE OF CONTENTS$']
    last_toc_idx = -1

    for i, para in enumerate(doc.paragraphs):
        sname = para.style.name
        text = para.text.strip()
        # 目录标题段落
        if _matches_any(text, TOC_TITLE_PATS):
            last_toc_idx = i
            continue
        # toc 样式段落（目录条目）
        if ('toc' in sname.lower() or '目录' in sname) and last_toc_idx >= 0:
            last_toc_idx = i

    if last_toc_idx >= 0:
        return last_toc_idx + 1  # 目录后第一个段落

    # 无目录：用第一个匹配标准章节模式的段落作为正文起点
    # 防止封面/声明页中偶发的 Heading 样式段落被当作章节处理
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() and _matches_any(para.text.strip(), HEADING1_PATTERNS):
            return i

    return 0  # 完全找不到任何章节标记，从头开始


def _insert_chapter_section_breaks(doc, chapter_style):
    """
    在每个一级标题段落前插入 nextPage 分节符，使每章成为独立的节。
    同时移除章节标题上的 pageBreakBefore（分节符已保证换页）。
    处理所有 Heading 1 段落（含无编号章节如绪论、参考文献等）。
    返回 {section_idx: chapter_title_text} 字典（节索引 → 章节标题文字）。
    """
    # 只收集目录之后的章节标题段落（避免封面/摘要等区域被误识别为章节）
    post_toc_start = _find_post_toc_idx(doc)
    chapter_entries = []
    for i, para in enumerate(doc.paragraphs):
        if i < post_toc_start:
            continue
        if para.style.name.lower() == chapter_style.lower() and para.text.strip():
            chapter_entries.append((i, para.text.strip()))

    if not chapter_entries:
        return {}

    body = doc.element.body
    body_sectPr = body.find(qn('w:sectPr'))

    for entry_idx, (idx, _chapter_text) in enumerate(chapter_entries):
        # 所有章节（含第一章）都在其前一段落插入分节符
        # 第一章：在目录最后一段与第一章之间插入分节符
        if idx == 0:
            continue  # 文档第一段落无需处理
        prev_para = doc.paragraphs[idx - 1]
        pPr = prev_para._p.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            continue  # 已有分节符

        new_sectPr = OxmlElement('w:sectPr')
        pg_type = OxmlElement('w:type')
        pg_type.set(qn('w:val'), 'nextPage')
        new_sectPr.append(pg_type)
        if body_sectPr is not None:
            for tag in ('w:pgSz', 'w:pgMar', 'w:cols', 'w:docGrid'):
                elem = body_sectPr.find(qn(tag))
                if elem is not None:
                    new_sectPr.append(copy.deepcopy(elem))

        if pPr is None:
            pPr = OxmlElement('w:pPr')
            prev_para._p.insert(0, pPr)
        pPr.append(new_sectPr)

        # 分节符已换页，移除章节标题的 pageBreakBefore 避免产生多余空白页
        ch_pPr = doc.paragraphs[idx]._p.find(qn('w:pPr'))
        if ch_pPr is not None:
            pgBr = ch_pPr.find(qn('w:pageBreakBefore'))
            if pgBr is not None:
                ch_pPr.remove(pgBr)

    # 扫描 body XML 构建 section_idx → 章节标题 映射
    style_id_map = {}
    for style in doc.styles:
        sid = style.element.get(qn('w:styleId'), '')
        style_id_map[sid] = style.name

    section_idx = 0
    chapter_section_map = {}

    for elem in body:
        local = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if local == 'p':
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    sid = pStyle.get(qn('w:val'), '')
                    sname = style_id_map.get(sid, '')
                    if sname.lower() == chapter_style.lower():
                        runs_text = ''.join(
                            t.text or '' for t in elem.findall('.//' + qn('w:t'))
                        ).strip()
                        if runs_text:
                            if section_idx not in chapter_section_map:
                                chapter_section_map[section_idx] = runs_text
                inner_sectPr = pPr.find(qn('w:sectPr'))
                if inner_sectPr is not None:
                    section_idx += 1

    # 正文第一章从第1页开始计算（前置页码独立，正文重新从1开始）
    if chapter_section_map:
        ch1_section_idx = min(chapter_section_map.keys())
        try:
            doc_sections = list(doc.sections)
            if ch1_section_idx < len(doc_sections):
                target_sectPr = doc_sections[ch1_section_idx]._sectPr
                pg_num_type = target_sectPr.find(qn('w:pgNumType'))
                if pg_num_type is None:
                    pg_num_type = OxmlElement('w:pgNumType')
                    target_sectPr.append(pg_num_type)
                pg_num_type.set(qn('w:start'), '1')
        except Exception:
            pass

    return chapter_section_map


def _remove_all_header_references(doc):
    """
    从文档所有 sectPr（含嵌入式分节符和 body 末尾）中移除 w:headerReference。
    调用后所有节均变为"链接到上一节"状态，可干净地重新建立页眉。
    """
    body = doc.element.body
    # 嵌入式 sectPr（位于段落的 w:pPr 内）
    for p_elem in body.iter(qn('w:p')):
        pPr = p_elem.find(qn('w:pPr'))
        if pPr is not None:
            sp = pPr.find(qn('w:sectPr'))
            if sp is not None:
                for ref in list(sp.findall(qn('w:headerReference'))):
                    sp.remove(ref)
    # body 末尾 sectPr
    body_sp = body.find(qn('w:sectPr'))
    if body_sp is not None:
        for ref in list(body_sp.findall(qn('w:headerReference'))):
            body_sp.remove(ref)


def _set_header_paragraph(header, doc, text=None, use_styleref=False,
                          styleref_style='Heading 1', backward_lookup=True):
    """
    设置页眉段落内容。
    text=None + use_styleref=False → 空白页眉
    use_styleref=True → StyleRef 域（自动章节标题）
    text='xxx' → 静态文字
    """
    # 确保该节有独立的页眉（不继承上一节）
    try:
        if header.is_linked_to_previous:
            header.is_linked_to_previous = False
    except Exception:
        return  # 无法建立独立页眉，跳过

    # 清空已有内容
    _clear_header_content(header)

    # 获取段落：优先用已有段落，否则直接在 XML 中插入
    if header.paragraphs:
        para_elem = header.paragraphs[0]._p
    else:
        # _Header 无 add_paragraph()，直接操作底层 XML
        try:
            hdr_xml = header._element  # w:hdr 元素
        except Exception:
            hdr_xml = None
        if hdr_xml is None:
            return
        para_elem = OxmlElement('w:p')
        hdr_xml.append(para_elem)

    # 设置段落样式为 Header，并强制居中
    pPr = para_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para_elem.insert(0, pPr)
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = OxmlElement('w:pStyle')
        pPr.insert(0, pStyle)
    pStyle.set(qn('w:val'), 'Header')
    # 强制居中（覆盖 Header 样式可能继承的左对齐）
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'center')

    if use_styleref:
        for elem in _make_styleref_field(styleref_style, backward_lookup=backward_lookup):
            para_elem.append(elem)
    elif text:
        r = OxmlElement('w:r')
        r.append(_make_rPr())
        t = OxmlElement('w:t')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        para_elem.append(r)
    # else: 空白，不追加任何 run


def _get_sections_with_front_matter(doc):
    """
    返回包含"摘要目录标题"样式段落的 section 索引集合。
    这类 section 需要用 STYLEREF "摘要目录标题" 作为奇数页页眉。
    """
    body = doc.element.body
    section_idx = 0
    sections_with_fm = set()

    style_id_map = {}
    for style in doc.styles:
        sid = style.element.get(qn('w:styleId'), '')
        style_id_map[sid] = style.name

    for elem in body:
        local = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if local == 'p':
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    sid = pStyle.get(qn('w:val'), '')
                    if style_id_map.get(sid) == '摘要目录标题':
                        sections_with_fm.add(section_idx)
                inner_sectPr = pPr.find(qn('w:sectPr'))
                if inner_sectPr is not None:
                    section_idx += 1
        elif local == 'tbl':
            section_idx += 0  # 表格内分节符较罕见，此处跳过

    return sections_with_fm


def update_headers(doc, thesis_title=''):
    """
    设置页眉（符合厦大规范）：
    - 为每个章节（第X章）插入 nextPage 分节符，每章独立成节
    - 章节节：奇数页 = 章节标题文字（静态，避免 STYLEREF 兼容性问题）
    - 前置节（含摘要目录标题）：奇数页 = STYLEREF "摘要目录标题"
    - 封面/其他节：奇数页留空
    - 所有节偶数页 = 论文标题
    """
    changes = []
    display_title = thesis_title.strip() if thesis_title.strip() else (_extract_thesis_title(doc) or '论文标题')

    # 清除所有现存的 w:headerReference，确保页眉从零开始构建（避免残留旧 STYLEREF 内容）
    _remove_all_header_references(doc)

    chapter_style = _get_chapter_heading_style(doc)

    # 为每章插入分节符，获取 section_idx → chapter_title 映射
    chapter_section_map = _insert_chapter_section_breaks(doc, chapter_style)

    # 获取含摘要目录标题的节集合
    sections_with_fm = _get_sections_with_front_matter(doc)

    for idx, section in enumerate(doc.sections):
        try:
            sectPr = section._sectPr

            # 启用奇偶页不同页眉
            if sectPr.find(qn('w:evenAndOddHeaders')) is None:
                sectPr.append(OxmlElement('w:evenAndOddHeaders'))

            chapter_title = chapter_section_map.get(idx)

            # 奇数页页眉
            if chapter_title is not None:
                # 正文章节节：直接写入章节标题文字（不依赖 STYLEREF）
                _set_header_paragraph(section.header, doc, text=chapter_title)
            elif idx in sections_with_fm:
                # 前置（摘要/目录）：STYLEREF 摘要目录标题
                _set_header_paragraph(section.header, doc, use_styleref=True,
                                      styleref_style='摘要目录标题', backward_lookup=False)
            else:
                # 封面等：留空
                _set_header_paragraph(section.header, doc)

            # 偶数页页眉：论文标题
            _set_header_paragraph(section.even_page_header, doc, text=display_title)

        except Exception:
            pass

    chapter_count = len(chapter_section_map)
    fm_count = len(sections_with_fm)
    changes.append(
        f'页眉：检测到 {chapter_count} 个章节（各自独立成节），章节页=章节标题，'
        f'前置 {fm_count} 节=摘要目录标题，偶数页="{display_title}"'
    )
    return changes


def _make_toc_field_paragraphs(doc, skip_paras=None, translate=False, static=False):
    """
    创建 TOC 段落列表，并从文档标题预填充目录条目。
    static=False（默认）：中文目录模式，生成 TOC 域（Word 打开自动更新页码）。
    static=True：英文目录模式，生成静态段落，避免 Word 更新域时覆盖翻译文字。
    skip_paras: 应跳过（不作为标题来源）的段落集合。
    translate: True 时将条目文字翻译为英文。
    """
    skip_paras = skip_paras or set()

    # 使用 outlineLvl 检测所有标题样式（支持 Heading N、标题 N 及任意自定义名）
    doc_heading_styles = _get_heading_styles(doc)

    # 动态检测最高级标题，使其映射为 toc 1
    min_level = None
    for para in doc.paragraphs:
        if id(para) in skip_paras:
            continue
        lv = doc_heading_styles.get(para.style.name)
        if lv is not None and para.text.strip():
            if min_level is None or lv < min_level:
                min_level = lv
    if min_level is None:
        min_level = 1

    # 构建 heading 样式 → toc 样式 映射（相对级别）
    heading_to_toc = {}
    for sname, lvl in doc_heading_styles.items():
        rel = lvl - min_level + 1
        if 1 <= rel <= 3:
            heading_to_toc[sname] = f'toc {rel}'
    # 兼容性：确保标准 Heading N 名也在映射中
    for offset in range(3):
        sname = f'Heading {min_level + offset}'
        if sname not in heading_to_toc:
            heading_to_toc[sname] = f'toc {offset + 1}'
    toc_instr = f' TOC \\o "{min_level}-{min_level + 2}" \\h \\z \\u '

    # 从文档中收集标题条目
    entries = []  # (toc_style, display_text, source_para)
    for para in doc.paragraphs:
        if id(para) in skip_paras:
            continue
        toc_style = heading_to_toc.get(para.style.name)
        if toc_style and para.text.strip():
            entry_text = para.text.strip()
            if translate:
                entry_text = _translate_toc_entry(entry_text)
            entries.append((toc_style, entry_text, para))

    # 静态英文目录：为每个标题段落添加书签，以便用 PAGEREF 引用页码
    if static and entries:
        bm_base = _get_max_bookmark_id(doc) + 1
        bm_names = []
        for idx, (_, _, para) in enumerate(entries):
            bm_name = f'_xmutoc{bm_base + idx}'
            _add_bookmark(para._p, bm_base + idx, bm_name)
            bm_names.append(bm_name)
    else:
        bm_names = [None] * len(entries)

    result_paras = []

    if not static:
        # ── 动态模式：TOC 域（中文目录用，Word 打开自动更新）──
        p_start = OxmlElement('w:p')
        r1 = OxmlElement('w:r')
        fc1 = OxmlElement('w:fldChar')
        fc1.set(qn('w:fldCharType'), 'begin')
        fc1.set(qn('w:dirty'), '1')
        r1.append(fc1)
        p_start.append(r1)

        r2 = OxmlElement('w:r')
        instr = OxmlElement('w:instrText')
        instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instr.text = toc_instr
        r2.append(instr)
        p_start.append(r2)

        r3 = OxmlElement('w:r')
        fc2 = OxmlElement('w:fldChar')
        fc2.set(qn('w:fldCharType'), 'separate')
        r3.append(fc2)
        p_start.append(r3)

        if not entries:
            r_ph = OxmlElement('w:r')
            t_ph = OxmlElement('w:t')
            t_ph.text = '（在 Word 中按 F9 更新目录）'
            r_ph.append(t_ph)
            p_start.append(r_ph)
            r_end = OxmlElement('w:r')
            fc_end = OxmlElement('w:fldChar')
            fc_end.set(qn('w:fldCharType'), 'end')
            r_end.append(fc_end)
            p_start.append(r_end)
            result_paras.append(p_start)
            return result_paras

        result_paras.append(p_start)
    elif not entries:
        # 静态模式但无条目：返回空
        return result_paras

    # ── 每个标题一段（动态模式作为域内容，静态模式作为独立段落）──
    for idx, (toc_style_name, text, _para) in enumerate(entries):
        p = OxmlElement('w:p')

        # 段落样式
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        try:
            style_elem = doc.styles[toc_style_name].element
            sid = style_elem.get(qn('w:styleId'), toc_style_name.replace(' ', ''))
        except KeyError:
            sid = toc_style_name.replace(' ', '')
        pStyle.set(qn('w:val'), sid)
        pPr.append(pStyle)

        # 静态英文目录：按层级添加直接缩进（toc 1=0, toc 2=360, toc 3=720 twips）
        if static:
            _EN_TOC_INDENT = {'toc 1': 0, 'toc 2': 360, 'toc 3': 720}
            left_twips = _EN_TOC_INDENT.get(toc_style_name, 0)
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(left_twips))
            ind.set(qn('w:firstLine'), '0')
            pPr.append(ind)

        p.append(pPr)

        # 标题文字（去除前导空白）
        r_text = OxmlElement('w:r')
        t_elem = OxmlElement('w:t')
        t_elem.text = text.lstrip()
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r_text.append(t_elem)
        p.append(r_text)

        # Tab + 页码
        r_tab = OxmlElement('w:r')
        r_tab.append(OxmlElement('w:tab'))
        p.append(r_tab)

        bm_name = bm_names[idx] if idx < len(bm_names) else None
        if static and bm_name:
            # 静态英文目录：用 PAGEREF 域引用标题所在页码
            for pg_r in _make_pageref_runs(bm_name):
                p.append(pg_r)
        else:
            r_pg = OxmlElement('w:r')
            t_pg = OxmlElement('w:t')
            t_pg.text = '?'
            r_pg.append(t_pg)
            p.append(r_pg)

        # 动态模式：最后一条目附上 end fldChar
        if not static and idx == len(entries) - 1:
            r_end = OxmlElement('w:r')
            fc_end = OxmlElement('w:fldChar')
            fc_end.set(qn('w:fldCharType'), 'end')
            r_end.append(fc_end)
            p.append(r_end)

        result_paras.append(p)

    return result_paras


def _enable_update_fields(doc):
    """在文档设置中启用"打开时自动更新所有域"（含目录）"""
    try:
        settings_elem = doc.settings.element
        uf = settings_elem.find(qn('w:updateFields'))
        if uf is None:
            uf = OxmlElement('w:updateFields')
            settings_elem.append(uf)
        uf.set(qn('w:val'), '1')
    except Exception:
        pass


# ── 中英文学术词汇词典 ────────────────────────────────────────────────────────
_ZH_NUMS = {
    '一': '1', '二': '2', '三': '3', '四': '4', '五': '5',
    '六': '6', '七': '7', '八': '8', '九': '9', '十': '10',
    '十一': '11', '十二': '12', '十三': '13', '十四': '14', '十五': '15',
}

_ZH_VOCAB = {
    # 章节结构
    '绪论': 'Introduction', '结论': 'Conclusions', '展望': 'Prospects',
    '未来展望': 'Future Prospects', '总结': 'Summary', '小结': 'Summary',
    '本章小结': 'Chapter Summary', '结论与展望': 'Conclusions and Prospects',
    '参考文献': 'References', '致谢': 'Acknowledgments',
    '附录': 'Appendix', '摘要': 'Abstract',
    # 研究通用
    '研究背景': 'Research Background', '背景': 'Background',
    '问题提出': 'Problem Statement', '问题': 'Issues',
    '研究目的': 'Research Objectives', '研究目标': 'Research Objectives',
    '目的': 'Objectives', '目标': 'Objectives',
    '研究意义': 'Research Significance', '意义': 'Significance',
    '实践价值': 'Practical Value', '理论意义': 'Theoretical Significance',
    '创新点': 'Innovation Points', '创新': 'Innovation',
    '研究内容': 'Research Content', '主要内容': 'Main Content', '主要': 'Main',
    '内容': 'Content', '内容安排': 'Content Arrangement', '安排': 'Arrangement',
    '研究方法': 'Research Methods', '方法': 'Methods',
    '研究方法论': 'Research Methodology', '方法论': 'Methodology',
    '选择与依据': 'Selection and Basis',
    '技术路线': 'Technical Roadmap', '路线图': 'Roadmap',
    '论文结构': 'Thesis Structure', '结构': 'Structure',
    '阐述': 'Elaboration', '界定': 'Definition',
    '本研究': 'This Study', '本文': 'This Paper',
    '框架引出': 'Framework Introduction',
    # 理论
    '文献综述': 'Literature Review', '文献': 'Literature', '综述': 'Review',
    '相关理论': 'Related Theory', '理论基础': 'Theoretical Foundation',
    '理论': 'Theory', '基础': 'Foundation', '基础与': 'Foundations and',
    '理论框架': 'Theoretical Framework', '总体框架': 'Overall Framework',
    '运行逻辑': 'Operational Logic', '总体架构': 'Overall Architecture',
    '设计原则': 'Design Principles', '原则': 'Principles',
    '内涵': 'Connotation', '发展演进': 'Development and Evolution',
    '应用研究': 'Applied Research', '应用': 'Application',
    '述评': 'Review and Commentary', '适用性': 'Applicability',
    '批判': 'Critique', '局限': 'Limitations', '局限性': 'Limitations',
    '适用性批判与局限性分析': 'Applicability Critique and Limitation Analysis',
    '互补性': 'Complementarity', '缺口': 'Gap',
    '互补性与融合缺口分析': 'Complementarity and Integration Gap Analysis',
    '内在逻辑': 'Internal Logic', '可行性': 'Feasibility',
    '必要性': 'Necessity', '必要性与可行性分析': 'Necessity and Feasibility Analysis',
    '融合': 'Integration', '融合模型': 'Integration Model',
    '理论融合': 'Theoretical Integration',
    # 公司/行业
    '营销': 'Marketing', '营销环境': 'Marketing Environment',
    '行业': 'Industry', '行业环境': 'Industry Environment',
    '行业分析': 'Industry Analysis',
    '宏观环境': 'Macro Environment',
    '竞争格局': 'Competitive Landscape',
    '市场需求': 'Market Demand', '市场': 'Market',
    '公司': 'Company', '企业': 'Enterprise',
    '现状': 'Current Situation', '现状介绍': 'Current Situation Overview',
    '现状及': 'Situation and', '经营现状': 'Operational Status',
    '战略定位': 'Strategic Positioning', '战略': 'Strategy',
    '发展愿景': 'Development Vision', '愿景': 'Vision',
    '商业模式': 'Business Model', '服务体系': 'Service System',
    '核心产品': 'Core Products', '产品服务体系': 'Product and Service System',
    '解析': 'Analysis', '诊断': 'Diagnosis', '问题诊断': 'Problem Diagnosis',
    '增长瓶颈': 'Growth Bottleneck', '增长': 'Growth',
    '价值': 'Value', '服务': 'Service', '思想': 'Philosophy',
    '体系': 'System', '双模': 'Dual-Mode',
    '建设': 'Construction', '完善': 'Improvement', '提升': 'Enhancement',
    '设计': 'Design', '实现': 'Implementation', '探索': 'Exploration',
    '保障机制': 'Guarantee Mechanisms', '机制': 'Mechanisms',
    '组织架构': 'Organizational Structure',
    '团队协作': 'Team Collaboration', '调整': 'Adjustment',
    '数据基础设施': 'Data Infrastructure', '度量指标': 'Measurement Indicators',
    '指标体系': 'Indicator System',
    # 精益创业 / 模型
    '客户关系管理': 'Customer Relationship Management',
    '精益创业': 'Lean Startup', '精益': 'Lean',
    '创新管理': 'Innovation Management',
    '用户增长': 'User Growth', '用户': 'Users', '用户价值': 'User Value',
    '流量获取': 'Traffic Acquisition', '流量': 'Traffic',
    '非对称性': 'Asymmetry', '准入机制': 'Access Mechanism', '缺失': 'Absence',
    '价值主张': 'Value Proposition', '偏离': 'Deviation',
    '验证路径': 'Verification Pathway',
    '增长引擎': 'Growth Engine', '外部依赖': 'External Dependence',
    '不可持续性': 'Unsustainability',
    '资源错配': 'Resource Misallocation',
    '工作流程': 'Work Process', '工作': 'Work', '流程': 'Process',
    '认知滞后': 'Cognitive Lag',
    '最小可行性产品': 'Minimum Viable Product',
    '闭环验证': 'Closed-Loop Verification',
    '测量偏离': 'Measurement Deviation',
    '虚荣指标': 'Vanity Metrics', '核算单位经济效益': 'Unit Economics',
    '学习断裂': 'Learning Disconnect', '实验设计': 'Experimental Design',
    '系统性': 'Systematic',
    '核心策略': 'Core Strategies', '核心': 'Core',
    '策略': 'Strategy', '模型': 'Model',
    '动态准入': 'Dynamic Access', '准入': 'Access',
    '价值增长': 'Value-Growth', '双模验证': 'Dual-Mode Verification',
    '协同逻辑': 'Synergy Logic', '关键增长节点': 'Key Growth Nodes',
    '流量纯化': 'Traffic Purification', '算法反馈': 'Algorithm Feedback',
    '算法反向纠偏': 'Algorithm Reverse Correction',
    # 实施
    '实施': 'Implementation', '实施背景': 'Implementation Background',
    '实施目标': 'Implementation Objectives', '实施总结': 'Implementation Summary',
    '实施前后': 'Before and After Implementation',
    '关键数据': 'Key Data', '对比分析': 'Comparative Analysis',
    '实验背景': 'Experiment Background',
    '深挖': 'In-depth Exploration',
    '实施过程': 'Implementation Process', '迭代': 'Iteration',
    '资源优化配置': 'Resource Optimization and Allocation',
    '反思': 'Reflection', '实施总结与反思': 'Implementation Summary and Reflection',
    '具体应用案例详述': 'Detailed Case Application',
    '量化设定': 'Quantitative Setting',
    '效果评估': 'Effect Evaluation',
    '规模驱动': 'Scale-Driven', '价值驱动': 'Value-Driven',
    '战略转型': 'Strategic Transformation',
    # 数据/PEST
    'PEST分析': 'PEST Analysis', '分析': 'Analysis',
    '模型理论及述评': 'Model Theory and Commentary',
    '理论内涵与发展演进': 'Theoretical Connotation and Development',
    '传统互联网领域': 'Traditional Internet Field',
    '核心框架与原则': 'Core Framework and Principles',
    '企业创新管理': 'Enterprise Innovation Management',
    # 通用动词/形容词/名词
    '发展': 'Development', '特征': 'Characteristics',
    '路径': 'Pathway', '缺乏': 'Lack of', '过度关注': 'Over-focus on',
    '而忽视': 'while Ignoring', '开展': 'Conduct',
    '介绍': 'Overview', '提出': 'Proposal',
    '选择': 'Selection', '依据': 'Basis',
    '调整': 'Adjustment', '构建': 'Construction',
    '研究': 'Research', '结论': 'Conclusions',
    '研究结论': 'Research Conclusions',
    '面临': 'Facing', '挑战': 'Challenges',
    '初创企业': 'Startup Companies', '初创': 'Startup',
    '现有': 'Existing', '环境': 'Environment',
    '系统': 'System', '测量': 'Measurement', '认知': 'Cognitive',
    '学习': 'Learning', '断裂': 'Disconnect',
    '滞后': 'Lag', '偏离': 'Deviation',
    # 补充单词
    '案例': 'Case', '阶段': 'Phase', '步骤': 'Step',
    '网关': 'Gateway', '清洗': 'Cleansing',
    '数据': 'Data', '纯化': 'Purification',
    '实验': 'Experiment', '过程': 'Process',
    '资源': 'Resources', '优化': 'Optimization', '配置': 'Allocation',
    '架构': 'Architecture', '设施': 'Infrastructure',
    '指标': 'Metrics', '度量': 'Measurement', '量化': 'Quantification',
    '深挖': 'In-depth Exploration',
    '比较': 'Comparison', '对比': 'Comparison',
    # 通用连接词（转义为空格，避免拼接错误）
    '与': ' and ', '及': ' and ', '和': ' and ', '在': ' in ',
    '中的': ' in ', '中': ' in ', '的': ' ', '下的': ' under ',
    '基于': 'Based on ', '对': ' on ', '向': ' toward ',
    '从': 'from ', '为': ' for ', '走向': 'toward ',
    # 顿号 → 逗号
    '、': ', ',
    # 全角标点（保留弯引号以便后处理区分开/闭）
    '：': ': ', '"': '\u201c', '"': '\u201d', '（': '(', '）': ')',
    '——': '—', '…': '...',
}


def _translate_toc_entry(text):
    """
    将中文目录条目翻译为英文。
    处理"第X章"、"第X节"、"一、"等结构，剩余部分查词典翻译。
    """
    text = text.strip()
    if not text:
        return text

    # 第X章 Y → Chapter N  Y_en
    m = re.match(r'^第([\u4e00-\u9fa5]+|\d+)\s*章\s*(.*)', text)
    if m:
        num = _ZH_NUMS.get(m.group(1), m.group(1))
        rest = _translate_phrase(m.group(2).strip(': ：'))
        return f'Chapter {num}  {rest}'.strip()

    # 第X节：Y → Section N: Y_en
    m = re.match(r'^第([\u4e00-\u9fa5]+|\d+)\s*节[：:]\s*(.*)', text)
    if m:
        num = _ZH_NUMS.get(m.group(1), m.group(1))
        rest = _translate_phrase(m.group(2).strip())
        return f'Section {num}: {rest}'.strip()

    # 一、二、三、 Y → 1. Y_en
    m = re.match(r'^([\u4e00-\u9fa5]+)[、．.]\s*(.*)', text)
    if m and m.group(1) in _ZH_NUMS:
        num = _ZH_NUMS[m.group(1)]
        rest = _translate_phrase(m.group(2).strip())
        return f'{num}. {rest}'.strip()

    # 参考文献
    if re.match(r'^参考文献', text):
        return 'References'

    return _translate_phrase(text)


def _translate_phrase(text):
    """用词典对短语做最长匹配翻译，未识别部分保留原文。"""
    if not text:
        return text
    if text in _ZH_VOCAB:
        return _ZH_VOCAB[text].strip()

    result = []
    i = 0
    while i < len(text):
        matched = False
        for end in range(min(i + 20, len(text)), i, -1):
            seg = text[i:end]
            if seg in _ZH_VOCAB:
                # 词典匹配：前后加空格确保单词边界
                result.append(' ')
                result.append(_ZH_VOCAB[seg])
                result.append(' ')
                i = end
                matched = True
                break
        if not matched:
            result.append(text[i])
            i += 1

    # 合并并清理多余空格，去除标点前的空格
    translated = re.sub(r' +', ' ', ''.join(result)).strip()
    translated = re.sub(r' ([,;:!?])', r'\1', translated)
    # 去除引号内侧多余空格（弯引号区分开/闭）
    translated = re.sub(u'\u201c\\s+', u'\u201c', translated)   # 左引号后不空格
    translated = re.sub(u'\\s+\u201d', u'\u201d', translated)   # 右引号前不空格

    # 将残余的中文数字单字转为阿拉伯数字（如"实验一" → "Experiment 1"）
    def _replace_zh_num(m):
        return _ZH_NUMS.get(m.group(0), m.group(0))
    translated = re.sub(r'[一二三四五六七八九十]', _replace_zh_num, translated)

    return translated


def _make_toc_heading_para(doc, title_text, page_break=True):
    """
    创建目录标题段落（"目　录" 或 "Contents"），使用摘要目录标题样式。
    page_break=True 时在段落前自动分页。
    """
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    try:
        sid = doc.styles['摘要目录标题'].element.get(qn('w:styleId'), '摘要目录标题')
    except KeyError:
        sid = 'Normal'
    pStyle.set(qn('w:val'), sid)
    pPr.append(pStyle)
    if page_break:
        pgBr = OxmlElement('w:pageBreakBefore')
        pgBr.set(qn('w:val'), '1')
        pPr.append(pgBr)
    p.append(pPr)
    r = OxmlElement('w:r')
    r.append(_make_rPr(east_asia_font='黑体', size_pt=16))
    t = OxmlElement('w:t')
    t.text = title_text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p


def regenerate_toc(doc):
    """
    重新生成中英文目录。
    - 若已有"目录"/"Contents"标题段落：替换其后旧条目，插入 TOC 域。
    - 若均不存在：在第一章前自动插入"目　录"和"Contents"两个目录节。
    Word 打开时自动更新，或按 Ctrl+A → F9 手动更新。
    """
    changes = []

    toc_style_names = {
        'toc 1', 'toc 2', 'toc 3', 'toc 4', 'toc 5',
        'TOC 1', 'TOC 2', 'TOC 3', 'TOC 4', 'TOC 5',
        '目录1', '目录2', '目录3', '目录 1', '目录 2', '目录 3',
    }
    zh_toc_patterns = [r'^目\s*录$', r'^目　录$']
    en_toc_patterns = [r'^Contents?$', r'^TABLE OF CONTENTS$']

    def _has_toc_field(para):
        for r in para._p.findall('.//' + qn('w:instrText')):
            if r.text and 'TOC' in r.text.upper():
                return True
        return False

    all_paras = list(doc.paragraphs)
    found_zh = False
    found_en = False
    i = 0

    while i < len(all_paras):
        para = all_paras[i]
        text = para.text.strip()
        is_zh = _matches_any(text, zh_toc_patterns)
        is_en = _matches_any(text, en_toc_patterns)

        if (is_zh or is_en) and para.style.name not in {'Normal', '正文', 'Body Text'}:
            lang = '中文' if is_zh else '英文'
            if is_zh:
                found_zh = True
            else:
                found_en = True
            heading_p_elem = para._p

            if i + 1 < len(all_paras) and _has_toc_field(all_paras[i + 1]):
                changes.append(f'{lang}目录已含 TOC 域，跳过重新生成')
                i += 1
                continue

            # 收集旧目录条目
            toc_entry_paras = []
            j = i + 1
            while j < len(all_paras):
                nxt = all_paras[j]
                if nxt.style.name in toc_style_names:
                    toc_entry_paras.append(nxt)
                    j += 1
                elif (not nxt.text.strip() and j + 1 < len(all_paras)
                      and all_paras[j + 1].style.name in toc_style_names):
                    toc_entry_paras.append(nxt)
                    j += 1
                else:
                    break

            for toc_para in toc_entry_paras:
                p_elem = toc_para._p
                parent = p_elem.getparent()
                if parent is not None:
                    parent.remove(p_elem)

            toc_ps = _make_toc_field_paragraphs(
                doc, skip_paras={id(para)},
                translate=is_en, static=is_en)
            for tp in reversed(toc_ps):
                heading_p_elem.addnext(tp)

            entry_count = len(toc_ps)
            changes.append(
                f'重新生成{lang}目录：删除旧条目 {len(toc_entry_paras)} 行，'
                f'已插入 {entry_count} 条目录条目（页码待更新）'
            )
            i = j
        else:
            i += 1

    # ── 若有中文但无英文（或两者都没有）：插入缺失的目录节 ──
    if not found_zh or not found_en:
        chapter_style = _get_chapter_heading_style(doc)
        post_toc = _find_post_toc_idx(doc)
        first_chapter_para = None
        for i, para in enumerate(doc.paragraphs):
            if i < post_toc:
                continue
            if para.style.name == chapter_style and para.text.strip():
                first_chapter_para = para
                break

        if first_chapter_para is not None:
            skip_ids = {id(first_chapter_para)}
            anchor = first_chapter_para._p
            entry_count = 0

            if not found_zh and not found_en:
                # 两者都没有：先插英文（在anchor前），再插中文（在英文前）
                # 最终顺序：目　录 → 中文条目 → Contents → 英文条目 → 第一章

                # 1. 先把 Contents 插到 anchor 前
                p_en_title = _make_toc_heading_para(doc, 'Contents', page_break=True)
                toc_en_ps = _make_toc_field_paragraphs(doc, skip_paras=skip_ids, translate=True, static=True)
                anchor.addprevious(p_en_title)
                for tp in toc_en_ps:
                    anchor.addprevious(tp)

                # 2. 再把 目　录 插到 p_en_title 前
                p_zh_title = _make_toc_heading_para(doc, '目　录', page_break=True)
                toc_zh_ps = _make_toc_field_paragraphs(doc, skip_paras=skip_ids)
                p_en_title.addprevious(p_zh_title)
                for tp in toc_zh_ps:
                    p_en_title.addprevious(tp)

                entry_count = max(0, len(toc_zh_ps) - 1)
                changes.append(
                    f'自动插入中文目录（目　录）：预填充 {entry_count} 条目录条目'
                )
                entry_count = max(0, len(toc_en_ps) - 1)
                changes.append(
                    f'自动插入英文目录（Contents）：预填充 {entry_count} 条目录条目'
                )
                changes.append('页码显示为"?"，在 Word 中按 Ctrl+A → F9 更新')

            elif found_zh and not found_en:
                # 只有中文，补插英文：插到第一章前
                p_en_title = _make_toc_heading_para(doc, 'Contents', page_break=True)
                toc_en_ps = _make_toc_field_paragraphs(doc, skip_paras=skip_ids, translate=True, static=True)
                anchor.addprevious(p_en_title)
                for tp in toc_en_ps:
                    anchor.addprevious(tp)
                entry_count = max(0, len(toc_en_ps) - 1)
                changes.append(
                    f'自动插入英文目录（Contents）：预填充 {entry_count} 条目录条目'
                )

            elif found_en and not found_zh:
                # 只有英文，补插中文：插到第一章前（英文已在后面）
                p_zh_title = _make_toc_heading_para(doc, '目　录', page_break=True)
                toc_zh_ps = _make_toc_field_paragraphs(doc, skip_paras=skip_ids)
                anchor.addprevious(p_zh_title)
                for tp in toc_zh_ps:
                    anchor.addprevious(tp)
                entry_count = max(0, len(toc_zh_ps) - 1)
                changes.append(
                    f'自动插入中文目录（目　录）：预填充 {entry_count} 条目录条目'
                )
        else:
            if not found_zh and not found_en:
                changes.append('未找到章节标题段落，无法插入目录，已跳过')

    if found_zh or found_en or changes:
        _enable_update_fields(doc)

    return changes


def format_thesis(input_path, output_path, options=None):
    """
    主入口：对上传的论文文档应用厦大格式规范

    Args:
        input_path: 输入文档路径
        output_path: 输出文档路径
        options: 格式选项字典，可选键：
            - fix_page_setup (bool): 修正页面设置，默认True
            - fix_styles (bool): 修正样式定义，默认True
            - fix_heading_direct (bool): 清除标题直接格式，默认True
            - auto_detect_headings (bool): 自动识别标题，默认False
            - fix_body_fonts (bool): 修正正文字体，默认True
            - add_page_numbers (bool): 添加页码，默认True
            - add_headers (bool): 添加页眉，默认True
            - thesis_title (str): 论文标题（用于偶数页页眉）

    Returns:
        list: 所有修改内容的描述列表
    """
    if options is None:
        options = {}

    fix_page     = options.get('fix_page_setup', True)
    fix_styles   = options.get('fix_styles', True)
    fix_heading  = options.get('fix_heading_direct', True)
    auto_detect  = options.get('auto_detect_headings', False)
    fix_fonts    = options.get('fix_body_fonts', True)
    add_pgnum    = options.get('add_page_numbers', True)
    add_hdrs     = options.get('add_headers', True)
    regen_toc    = options.get('regenerate_toc', False)
    thesis_title = options.get('thesis_title', '')

    doc = Document(input_path)
    all_changes = []

    if fix_page:
        all_changes.extend(update_page_setup(doc))

    if fix_styles:
        all_changes.extend(update_styles(doc))

    if fix_heading:
        all_changes.extend(fix_heading_direct_format(doc))

    if auto_detect:
        all_changes.extend(auto_detect_headings(doc))

    if fix_fonts:
        all_changes.extend(fix_body_text_fonts(doc))

    # 题注格式 + 图表跨页防护（始终执行，只要 fix_styles 或 fix_fonts 之一开启）
    if fix_styles or fix_fonts:
        all_changes.extend(fix_captions_and_tables(doc))

    if regen_toc:
        all_changes.extend(regenerate_toc(doc))

    if add_hdrs:
        all_changes.extend(update_headers(doc, thesis_title))

    if add_pgnum:
        all_changes.extend(update_footer_page_numbers(doc))

    doc.save(output_path)
    return all_changes
